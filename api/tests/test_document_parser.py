from io import BytesIO

import fitz
from docx import Document

from app.application.document_parser import DocumentParseError, parse_document


def test_parse_markdown_removes_markup() -> None:
    parsed = parse_document("notes.md", b"# Title\n\nHello **Avpilot**.")

    assert "Title" in parsed.text
    assert "Hello Avpilot" in parsed.text
    assert "**" not in parsed.text


def test_parse_docx_includes_paragraphs_and_tables() -> None:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("实验记录")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "温度"
    table.cell(0, 1).text = "25℃"
    document.save(stream)

    parsed = parse_document("record.docx", stream.getvalue())

    assert "实验记录" in parsed.text
    assert "温度 25℃" in parsed.text


def test_parse_pdf_preserves_page_number() -> None:
    document = fitz.open()
    first = document.new_page()
    first.insert_text((72, 72), "first page")
    second = document.new_page()
    second.insert_text((72, 72), "second page")
    content = document.tobytes()
    document.close()

    parsed = parse_document("paper.pdf", content)
    second_offset = parsed.text.index("second page")

    assert parsed.page_for_offset(0) == 1
    assert parsed.page_for_offset(second_offset) == 2


def test_rejects_unsupported_or_empty_document() -> None:
    try:
        parse_document("data.csv", b"a,b")
    except DocumentParseError as exc:
        assert "不支持" in str(exc)
    else:
        raise AssertionError("unsupported extension should fail")

    try:
        parse_document("empty.txt", b"   \n")
    except DocumentParseError as exc:
        assert "没有可检索文本" in str(exc)
    else:
        raise AssertionError("empty document should fail")
