"""Extract normalized text and page-aware source sections from supported documents."""

import io
from collections.abc import Iterable
from pathlib import Path

import chardet
import fitz
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from app.domain.documents import ParsedDocument, SourceSection

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt", ".html", ".htm"}


class DocumentParseError(ValueError):
    """The uploaded file cannot be converted into searchable text."""


def parse_document(file_name: str, content: bytes) -> ParsedDocument:
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentParseError(f"不支持 {extension or '未知'} 文件，可上传：{supported}")
    if extension == ".pdf":
        parts = _parse_pdf(content)
    elif extension == ".docx":
        parts = [(_parse_docx(content), None)]
    elif extension in {".md", ".markdown"}:
        parts = [(_parse_markdown(content), None)]
    elif extension in {".html", ".htm"}:
        parts = [(_parse_html(content), None)]
    else:
        parts = [(_decode_text(content), None)]
    parsed = _assemble(parts)
    if not parsed.text.strip():
        raise DocumentParseError("文档中没有可检索文本")
    return parsed


def _decode_text(content: bytes) -> str:
    detected = chardet.detect(content)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return content.decode(encoding)
    except (LookupError, UnicodeDecodeError):
        return content.decode("utf-8", errors="ignore")


def _parse_pdf(content: bytes) -> list[tuple[str, int | None]]:
    try:
        with fitz.open(stream=content, filetype="pdf") as document:
            return [(page.get_text("text"), page.number + 1) for page in document]
    except Exception as exc:
        raise DocumentParseError("PDF 文件损坏或无法解析") from exc


def _parse_docx(content: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("Word 文件损坏或无法解析") from exc
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            value = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if value:
                blocks.append(value)
    return "\n".join(blocks)


def _parse_markdown(content: bytes) -> str:
    html = markdown.markdown(_decode_text(content), extensions=["tables", "fenced_code"])
    return _html_to_text(html)


def _parse_html(content: bytes) -> str:
    return _html_to_text(_decode_text(content))


def _html_to_text(value: str) -> str:
    soup = BeautifulSoup(value, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(" ")


def _assemble(parts: Iterable[tuple[str, int | None]]) -> ParsedDocument:
    text_parts: list[str] = []
    sections: list[SourceSection] = []
    cursor = 0
    for raw_text, page in parts:
        text = _normalize(raw_text)
        if not text:
            continue
        if text_parts:
            text_parts.append("\n\n")
            cursor += 2
        start = cursor
        text_parts.append(text)
        cursor += len(text)
        sections.append(SourceSection(text=text, start_char=start, end_char=cursor, page=page))
    return ParsedDocument(text="".join(text_parts), sections=tuple(sections))


def _normalize(value: str) -> str:
    lines = [" ".join(line.split()) for line in value.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()
